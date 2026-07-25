from io import BytesIO

from docx import Document
from app.services.resume_parser import UnsupportedResumeType, extract_resume_text


def test_extracts_docx_resume_text() -> None:
    document = Document()
    document.add_paragraph("Jordan Lee")
    document.add_paragraph("Data Analyst | jordan.lee@example.com")
    document.add_paragraph("Skills: SQL, Python, Tableau, Excel")
    document.add_paragraph("Experience: Built weekly dashboards for a retail analytics team.")

    output = BytesIO()
    document.save(output)

    text = extract_resume_text("jordan-lee-resume.docx", output.getvalue())

    assert "Jordan Lee" in text
    assert "SQL, Python, Tableau, Excel" in text
    assert "weekly dashboards" in text


def test_extracts_rtf_resume_text() -> None:
    rtf_content = rb"{\rtf1\ansi\deff0 {\fonttbl{\f0 Arial;}}\f0\b Alex Smith\b0\par Software Engineer\par Skills: Python, SQL\par}"
    text = extract_resume_text("resume.rtf", rtf_content)
    assert "Alex Smith" in text
    assert "Software Engineer" in text


def test_rejects_unsupported_file_type() -> None:
    try:
        extract_resume_text("resume.txt", b"plain text")
    except UnsupportedResumeType as error:
        assert "PDF, DOCX, DOC, RTF, or ODT" in str(error)
    else:
        raise AssertionError("Expected UnsupportedResumeType")